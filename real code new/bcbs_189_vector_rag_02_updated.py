
import os
from dotenv import load_dotenv
from typing import List, Union
from pydantic import BaseModel
import pandas as pd
import re
import tiktoken

from langchain_community.document_loaders import PyPDFLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_openai import ChatOpenAI, OpenAIEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_core.prompts import PromptTemplate
from langchain_core.runnables import RunnableParallel, RunnablePassthrough, RunnableLambda
from langchain_core.output_parsers import StrOutputParser

from io import BytesIO
from PyPDF2 import PdfReader
from docx import Document

# --- Load environment ---
load_dotenv()

# --- Globals ---
rag_chain = None
llm_for_structured = None
model_name = "gpt-3.5-turbo"

# --- Format for RAG ---
def format_docs(docs):
    return "\n\n".join(doc.page_content for doc in docs)

rag_prompt = PromptTemplate(
    template="""
You are a Basel III regulatory assistant.
Answer ONLY using the context below.
If context is insufficient, say "Insufficient context."

{context}
Question: {question}
""",
    input_variables=["context", "question"]
)

# --- Schema Definitions ---
class BaselTestCase(BaseModel):
    test_title: str
    test_case_title: str
    test_case_description: str
    test_steps: List[str]
    test_data: List[str]
    expected_results: List[str]

class BaselTestCaseList(BaseModel):
    test_cases: List[BaselTestCase]

# --- Token Estimation ---
def estimate_tokens(text: str, model_name: str) -> int:
    try:
        enc = tiktoken.encoding_for_model(model_name)
    except KeyError:
        enc = tiktoken.get_encoding("cl100k_base")
    return len(enc.encode(text))

# --- Build Prompt ---
def build_testcase_prompt(rag_answer: str, max_cases: int = 2) -> str:
    return (
        "You are a Basel 3.1 compliance testing expert.\n"
        f"Based on the following document-derived context, generate {max_cases} test cases:\n\n"
        f"{rag_answer}\n\n"
        "Each test case must include:\n"
        "- Test title\n"
        "- Test case title and description\n"
        "- At least 3 test steps\n"
        "- Test data\n"
        "- Expected results\n\n"
        "Return only JSON under key 'test_cases'."
    )


# --- Load Files from Upload and Build Vector Store ---
def load_documents_from_uploads(uploaded_files: List[BytesIO], test_mode: bool = False):
    from langchain_core.documents import Document as LC_Document
    all_texts = []

    for file in uploaded_files:
        filename = file.name.lower()
        try:
            if filename.endswith(".pdf"):
                reader = PdfReader(file)
                for page in reader.pages:
                    text = page.extract_text()
                    if text:
                        all_texts.append(LC_Document(page_content=text.strip()))
            elif filename.endswith(".docx"):
                doc = Document(file)
                for para in doc.paragraphs:
                    if para.text.strip():
                        all_texts.append(LC_Document(page_content=para.text.strip()))
            elif filename.endswith(".txt"):
                content = file.read().decode("utf-8")
                for line in content.splitlines():
                    if line.strip():
                        all_texts.append(LC_Document(page_content=line.strip()))
        except Exception as e:
            print(f"⚠️ Failed to read {filename}: {e}")

    if test_mode:
        all_texts = all_texts[:2]
        for doc in all_texts:
            doc.page_content = doc.page_content[:1000]

    splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
    chunks = splitter.split_documents(all_texts)
    embeddings = OpenAIEmbeddings(model="text-embedding-3-small")
    vector_store = FAISS.from_documents(chunks, embeddings)
    return vector_store.as_retriever(search_type="similarity", search_kwargs={"k": 4})

# --- Init with Uploaded Files ---
def init_rag_chain(test_mode=False, uploaded_files: Union[List[BytesIO], None] = None):
    global rag_chain, llm_for_structured, model_name

    model_name = "gpt-3.5-turbo" if test_mode else "gpt-4"

    if uploaded_files:
        retriever = load_documents_from_uploads(uploaded_files, test_mode)
    else:
        raise ValueError("Uploaded files are required for initializing RAG.")

    rag_chain = (
        RunnableParallel({
            "context": retriever | RunnableLambda(format_docs),
            "question": RunnablePassthrough()
        })
        | rag_prompt
        | ChatOpenAI(model=model_name, temperature=0.2, max_tokens=512)
        | StrOutputParser()
    )

    llm_for_structured = ChatOpenAI(
        model=model_name,
        temperature=0,
        max_tokens=1024
    ).with_structured_output(BaselTestCaseList, method="function_calling")

    return model_name

# --- Main Test Case Function ---
def generate_test_cases_for_requirement(requirement: str):
    try:
        rag_answer = rag_chain.invoke(requirement)
        test_prompt = build_testcase_prompt(rag_answer)
        structured_result = llm_for_structured.invoke(test_prompt)

        rag_tokens = estimate_tokens(rag_answer, model_name)
        prompt_tokens = estimate_tokens(test_prompt, model_name)

        result_rows = []
        for tc in structured_result.test_cases:
            d = tc.dict()
            max_len = max(len(d['test_steps']), len(d['test_data']), len(d['expected_results']))
            for i in range(max_len):
                result_rows.append({
                    'requirement': requirement,
                    'test_title': d['test_title'],
                    'test_case_title': d['test_case_title'],
                    'test_case_description': d['test_case_description'],
                    'step_number': f"Step {i+1}",
                    'test_step': re.sub(r"^Step\s*\d+\s*:\s*", "", d['test_steps'][i]) if i < len(d['test_steps']) else "",
                    'test_data': d['test_data'][i] if i < len(d['test_data']) else "",
                    'expected_result': d['expected_results'][i] if i < len(d['expected_results']) else ""
                })

        return pd.DataFrame(result_rows), rag_tokens, prompt_tokens
    except Exception as e:
        return f"Error: {e}", 0, 0
